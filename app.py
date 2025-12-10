import streamlit as st
import dashscope
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import io
import re
import json
import requests  # 新增：用于调用 DeepSeek 等通用接口
from http import HTTPStatus

# ================= 配置区 =================
APP_TITLE = "AI 实验报告生成系统 (V11.0 旗舰版)"
APP_ICON = "🎓"

# 字体与颜色配置
STYLE_CONFIG = {
    "title_font": "黑体",
    "body_font_cn": "宋体",
    "body_font_en": "Times New Roman",
    "code_font": "Courier New",
    "answer_color": RGBColor(0, 0, 139),  # 深蓝色
    "code_color": RGBColor(50, 50, 50),   # 深灰色
    "code_bg_color": "F2F2F2"             # 浅灰背景
}

# ================= 1. 页面初始化 & CSS 优化 =================
st.set_page_config(
    page_title=APP_TITLE,
    page_icon=APP_ICON,
    layout="wide",
    initial_sidebar_state="expanded"
)

# 【CSS 黑科技】界面深度美化与汉化
st.markdown("""
<style>
    /* 汉化上传按钮 */
    .stFileUploader label { display: none; }
    .stFileUploader::after {
        content: "请将 .docx 实验报告文件拖拽到此处，或点击浏览";
        display: block;
        text-align: center;
        color: #666;
        padding: 10px;
        border: 2px dashed #ccc;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    /* 调整主标题样式 */
    h1 { color: #1E3A8A; }
    /* 侧边栏样式优化 */
    .css-1d391kg { background-color: #F8F9FA; }
</style>
""", unsafe_allow_html=True)

# ================= 2. 核心逻辑层 =================

def set_font_style(run, font_type="body", is_bold=False):
    """字体样式工厂"""
    if font_type == "code":
        run.font.name = STYLE_CONFIG["code_font"]
        run.font.size = Pt(10.5)
        run.font.color.rgb = STYLE_CONFIG["code_color"]
    else:
        run.font.name = STYLE_CONFIG["body_font_en"]
        run._element.rPr.rFonts.set(qn('w:eastAsia'), STYLE_CONFIG["body_font_cn"])
        run.font.size = Pt(12)
        run.font.color.rgb = STYLE_CONFIG["answer_color"]
    run.font.bold = is_bold

def add_shading(paragraph):
    """代码块背景渲染"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), STYLE_CONFIG["code_bg_color"]))
    paragraph._element.get_or_add_pPr().append(shading_elm)

def read_docx_content(file):
    """全文档扫描器"""
    doc = Document(file)
    full_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    return "\n".join(full_text)

def fill_report_core(source_file, ai_data):
    """
    【智能填空引擎 V11】：标题保护 + 废话定点清除
    """
    source_file.seek(0)
    doc = Document(source_file)
    
    # 关键词黑名单 (只删这些，绝不删标题)
    KEYWORD_MAP = {
        'code': ["直接粘贴过来", "格式为", "运行代码", "粘贴代码", "截图"],
        'analysis': ["具体分析", "改进的办法", "错误原因分析", "心得", "问题或错误"],
        'design': ["撰写上述内容", "基本语法", "上述内容"] 
    }
    
    filled_status = {'design': False, 'code': False, 'analysis': False}

    def write_content(cell, content, mode):
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.line_spacing = 1.25 
            
            if mode == 'code':
                # 识别题目序号 (1)
                if re.match(r'^\s*[（\(]\d+[）\)]', line):
                    run = p.add_run(line)
                    set_font_style(run, "body", is_bold=True)
                # 识别 SQL/代码关键字
                elif any(k in line.upper() for k in ["CREATE", "INSERT", "UPDATE", "SELECT", "ALTER", "DROP", "TABLE", "TRIGGER", "PROCEDURE", "DECLARE", "BEGIN", "END", ";", "--"]):
                    add_shading(p)
                    run = p.add_run(line)
                    set_font_style(run, "code")
                else:
                    run = p.add_run(line)
                    set_font_style(run, "body")
            else:
                run = p.add_run(line)
                set_font_style(run, "body")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "".join([p.text for p in cell.paragraphs])
                for key, keywords in KEYWORD_MAP.items():
                    if not filled_status[key]: 
                        if any(kw in cell_text for kw in keywords):
                            # 安全删除逻辑
                            for p in cell.paragraphs:
                                for kw in keywords:
                                    if kw in p.text:
                                        p.text = "" 
                            
                            write_content(cell, ai_data[key], mode=key)
                            filled_status[key] = True
                            break 

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def call_ai_engine(api_key, context, provider, model_name):
    """
    【通用 AI 调度中心】
    支持：阿里云 DashScope (Qwen), DeepSeek (OpenAI协议)
    """
    
    system_prompt = f"""
    你是一名计算机助教。请根据以下实验题目，生成满分实验报告。
    
    【格式要求 (XML)】：
    <design>简述实验工具及核心SQL语法。</design>
    <code>
    (1) [复述题目1]
    [SQL 代码]
    
    (2) [复述题目2]
    [SQL 代码]
    </code>
    <analysis>实验总结与心得。</analysis>

    【题目内容】：
    {context}
    """

    # --- 分支 1: 阿里云 DashScope ---
    if provider == "dashscope":
        dashscope.api_key = api_key
        try:
            response = dashscope.Generation.call(
                model=model_name, 
                prompt=system_prompt
            )
            if response.status_code == HTTPStatus.OK:
                return parse_ai_response(response.output.text)
            else:
                return {'success': False, 'error': f"阿里 API 报错: {response.message}"}
        except Exception as e:
            return {'success': False, 'error': f"系统错误: {str(e)}"}

    # --- 分支 2: DeepSeek (OpenAI 兼容协议) ---
    elif provider == "deepseek":
        url = "https://api.deepseek.com/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
        data = {
            "model": "deepseek-chat", # DeepSeek V3
            "messages": [
                {"role": "system", "content": "你是一个有用的助手。请按 XML 格式输出。"},
                {"role": "user", "content": system_prompt}
            ],
            "stream": False
        }
        try:
            resp = requests.post(url, headers=headers, json=data)
            if resp.status_code == 200:
                result = resp.json()
                content = result['choices'][0]['message']['content']
                return parse_ai_response(content)
            else:
                return {'success': False, 'error': f"DeepSeek 报错: {resp.text}"}
        except Exception as e:
            return {'success': False, 'error': f"网络请求错误: {str(e)}"}

def parse_ai_response(text):
    """统一解析 XML"""
    try:
        design = re.search(r'<design>(.*?)</design>', text, re.DOTALL)
        code = re.search(r'<code>(.*?)</code>', text, re.DOTALL)
        analysis = re.search(r'<analysis>(.*?)</analysis>', text, re.DOTALL)
        
        return {
            'success': True,
            'design': design.group(1).strip() if design else "AI生成格式缺失(Design)",
            'code': code.group(1).strip() if code else "AI生成格式缺失(Code)",
            'analysis': analysis.group(1).strip() if analysis else "AI生成格式缺失(Analysis)"
        }
    except:
        return {'success': False, 'error': "解析 AI 返回内容失败，请重试"}

# ================= 3. 界面交互层 =================

with st.sidebar:
    st.image("https://img.icons8.com/color/96/artificial-intelligence.png", width=60)
    st.title("控制面板")
    
    # --- 模型选择 ---
    st.markdown("### 🧠 模型设置")
    provider = st.selectbox(
        "选择 AI 厂商", 
        ("阿里云 (DashScope)", "DeepSeek (深度求索)"),
        index=0
    )
    
    if provider == "阿里云 (DashScope)":
        model_name = st.selectbox(
            "选择模型版本",
            ("qwen-plus (推荐, 均衡)", "qwen-turbo (免费额度多)", "qwen-max (最聪明)"),
            index=0
        )
        api_key_help = "https://bailian.console.aliyun.com/"
        provider_code = "dashscope"
        model_code = model_name.split(" ")[0]
        
    else: # DeepSeek
        model_name = "DeepSeek-V3"
        st.caption("DeepSeek 当前非常火爆，也是目前最便宜的 API 之一。")
        api_key_help = "https://platform.deepseek.com/api_keys"
        provider_code = "deepseek"
        model_code = "deepseek-chat"

    # --- API Key 输入 ---
    api_key = st.text_input("在此粘贴 API Key", type="password")

    # --- 获取指南 (手风琴) ---
    with st.expander("🎁 如何获取免费 API Key?"):
        st.markdown(f"""
        **1. 阿里云通义千问 (DashScope)**
        * 新用户开通通常送几百万 Token，几乎用不完。
        * [👉 点击领取 DashScope Key]({api_key_help})
        
        **2. DeepSeek (深度求索)**
        * 注册即送 10元 额度 (约 500万 Token)，非常良心。
        * [👉 点击注册 DeepSeek]({api_key_help})
        """)

    st.markdown("---")
    st.caption(f"当前引擎: {provider} - {model_code}")

# --- 主界面 ---
st.title(APP_TITLE)

step1 = st.container()
step2 = st.container()

with step1:
    st.markdown("### 1️⃣ 上传作业模板")
    uploaded_file = st.file_uploader("文件", type=['docx'])
    
    if uploaded_file:
        file_text = read_docx_content(uploaded_file)
        st.success(f"✅ 识别成功！文档字数: {len(file_text)}")
        
        # 使用 Expander 而不是 Tabs 以避免上下文报错
        with st.expander("👁️ 预览题目内容"):
            st.text(file_text)

with step2:
    if uploaded_file:
        st.markdown("### 2️⃣ 智能生成")
        
        if not api_key:
            st.warning("⚠️ 请先在左侧输入 API Key")
        else:
            if st.button("🚀 开始生成作业", type="primary", use_container_width=True):
                
                # 进度条
                progress_text = "AI 正在连接中..."
                my_bar = st.progress(0, text=progress_text)
                
                # 1. AI 推理
                my_bar.progress(30, text=f"🧠 正在调用 {model_name} 进行思考...")
                ai_result = call_ai_engine(api_key, file_text, provider_code, model_code)
                
                if ai_result['success']:
                    # 2. 写入文档
                    my_bar.progress(70, text="✍️ 正在执行定点替换与填空...")
                    final_docx = fill_report_core(uploaded_file, ai_result)
                    
                    my_bar.progress(100, text="🎉 完成！")
                    st.success("✨ 作业生成完毕！标题已保护，格式已优化。")
                    
                    # 3. 下载与预览
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        st.download_button(
                            label="📥 下载最终作业 (.docx)",
                            data=final_docx,
                            file_name=f"Done_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )
                    with col2:
                        with st.expander("查看生成的 SQL 代码"):
                            st.code(ai_result['code'], language='sql')
                else:
                    my_bar.empty()
                    st.error(f"❌ 生成失败: {ai_result['error']}")
    else:
        st.info("👆 请先在上方上传 .docx 文件")
